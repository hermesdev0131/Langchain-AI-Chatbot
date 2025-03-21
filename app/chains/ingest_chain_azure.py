import asyncio
import datetime
import logging
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import chardet
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import io
from app.config import settings

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_contents: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_contents))
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

async def initialize_ingest_chain(file_contents: bytes, filename: str, vector_store) -> dict:
    """
    Ingestion chain that processes an uploaded document by:
      1. Decoding the file contents (assumed UTF-8)
      2. Creating a Document with metadata
      3. Splitting the document into chunks
      4. Adding the chunks to the Azure AI Search index via the provided vector store

    Parameters:
      - file_contents: Binary content of the uploaded file.
      - filename: Name of the file.
      - vector_store: An initialized AzureSearch vector store instance.

    Returns:
      A dictionary indicating success.
    """
    try:
        if filename.lower().endswith(".pdf"):
            # Handle PDF separately
            text = extract_text_from_pdf(file_contents)
            if not text.strip():
                raise ValueError("Extracted text from PDF is empty.")
        elif filename.lower().endswith(".docx"):
            # Wrap the bytes in a BytesIO object
            docx_file = DocxDocument(io.BytesIO(file_contents))
            # Join all paragraphs together
            text = "\n".join([p.text for p in docx_file.paragraphs])
            if not text.strip():
                raise ValueError("Extracted text from DOCX is empty.")
        else:
            # Detect encoding for text files
            detection = chardet.detect(file_contents)
            encoding = detection.get("encoding")
            if not encoding:
                encoding = "utf-8"  # Default fallback encoding
            print("Detected encoding:", encoding)
            text = file_contents.decode(encoding)

    except Exception as e:
        logger.error("File decoding error: %s", e)
        raise Exception("File must be UTF-8 encoded text")
    
    # Create a Document object for the file
    doc = Document(
        page_content=text,
        metadata={
            "filename": filename,
            "timestamp": int(datetime.datetime.now().timestamp())
        }
    )
    
    # Split the document into manageable chunks
    splitter = RecursiveCharacterTextSplitter(
      separators=["\n\n", "\n", " ", ""],
      chunk_size=settings.CHUNK_SIZE,
      chunk_overlap=settings.CHUNK_OVERLAP
    )
    docs = splitter.split_documents([doc])
    print(f"Number of chunks: {len(docs)}")
    for i, d in enumerate(docs):
        print(f"Chunk {i} length: {len(d.page_content)}")

    # Ingest the chunks into the Azure AI Search vector store.
    # This call runs in a background thread since vector_store.add_documents is blocking.
    await asyncio.to_thread(vector_store.add_documents, docs)
    
    logger.info("File '%s' ingested successfully.", filename)
    return {"status": "success", "message": f"File '{filename}' ingested successfully."}
