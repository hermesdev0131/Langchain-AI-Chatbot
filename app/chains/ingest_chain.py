import asyncio
import datetime
import logging
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import chardet
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import io
from bs4 import BeautifulSoup
import requests
from app.config import settings
from functools import partial

logger = logging.getLogger(__name__)

def extract_text_from_url(url: str) -> str:
    try:
        response = requests.get(url)
        response.raise_for_status()
    except Exception as e:
        logger.error("Error fetching URL '%s': %s", url, e)
        raise Exception(f"Error fetching URL '{url}': {e}")
    
    html = response.text
    soup = BeautifulSoup(html, 'html.parser')
    text = soup.get_text(separator="\n")
    
    return text

def extract_text_from_pdf(file_contents: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_contents))
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""

    return text

async def initialize_ingest_chain_document(file_contents: bytes, filename: str, vector_store) -> dict:
    """
    Ingestion chain that processes an uploaded document by:
      1. Decoding the file contents (assumed UTF-8)
      2. Creating a Document with metadata
      3. Splitting the document into chunks
      4. Adding the chunks to the index via the provided vector store

    Parameters:
      - file_contents: Binary content of the uploaded file.
      - filename: Name of the file.
      - vector_store: An initialized vector store instance.

    Returns:
      A dictionary indicating success.
    """
    try:
        if filename.lower().endswith(".pdf"):
            # Handle PDF separately
            text = extract_text_from_pdf(file_contents)
            if not text.strip():
                raise ValueError("Extracted text from PDF is empty.")
        elif filename.lower().endswith(".docx"):
            # Wrap the bytes in a BytesIO object
            docx_file = DocxDocument(io.BytesIO(file_contents))
            # Join all paragraphs together
            text = "\n".join([p.text for p in docx_file.paragraphs])
            if not text.strip():
                raise ValueError("Extracted text from DOCX is empty.")
        else:
            # Detect encoding for text files
            detection = chardet.detect(file_contents)
            encoding = detection.get("encoding")
            if not encoding:
                encoding = "utf-8"  # Default fallback encoding
            print("Detected encoding:", encoding)
            text = file_contents.decode(encoding)

    except Exception as e:
        logger.error("File decoding error: %s", e)
        raise Exception("File must be UTF-8 encoded text")
    
    # Create a Document object for the file
    doc = Document(
        page_content=text,
        metadata={
            "filename": filename,
            "file_path": filename,
            "timestamp": int(datetime.datetime.now().timestamp())
        }
    )
    
    # Split the document into manageable chunks
    splitter = RecursiveCharacterTextSplitter(
      separators=["\n\n", "\n", " ", ""],
      chunk_size=settings.CHUNK_SIZE,
      chunk_overlap=settings.CHUNK_OVERLAP
    )
    docs = splitter.split_documents([doc])
    print(f"Number of chunks: {len(docs)}")
    for i, d in enumerate(docs):
        print(f"Chunk {i} length: {len(d.page_content)}")

    # Ingest the chunks into the vector store.
    # This call runs in a background thread since vector_store.add_documents is blocking.
    await asyncio.to_thread(vector_store.add_documents, docs)
    
    logger.info("File '%s' ingested successfully.", filename)
    return {"status": "success", "message": f"File '{filename}' ingested successfully."}

async def initialize_ingest_chain_url(url: str, vector_store) -> dict:
    """
    Ingestion chain for processing a URL.
    This function fetches the URL, extracts text content, splits it, and adds the chunks to the vector store.
    
    Parameters:
      - url: The URL to ingest.
      - vector_store: An initialized vector store instance.
      
    Returns:
      A dictionary indicating the ingestion status.
    """
    try:
        text = extract_text_from_url(url)
        if not text.strip():
            raise ValueError("Extracted text from URL is empty.")
    except Exception as e:
        logger.error("URL ingestion error for '%s': %s", url, e)
        raise Exception(f"URL ingestion error: {e}")
    
    doc = Document(
        page_content=text,
        metadata={
            "url": url,
            "timestamp": int(datetime.datetime.now().timestamp())
        }
    )
    
    splitter = RecursiveCharacterTextSplitter(
        separators=["\n\n", "\n", " ", ""],
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP
    )
    docs = splitter.split_documents([doc])
    logger.info("Number of chunks from URL '%s': %d", url, len(docs))
    
    await asyncio.to_thread(vector_store.add_documents, docs)
    
    logger.info("URL '%s' ingested successfully.", url)
    return {"status": "success", "message": f"URL '{url}' ingested successfully."}

# --- Ingestion Chain Wrapper Implementation ---

class IngestionChainWrapper:
    """
    Wrapper that holds ingestion functions for both documents and URLs.
    The functions are pre-bound with the vector_store dependency.
    """
    def __init__(self, vector_store):
        self.vector_store = vector_store
        # Pre-bind vector_store to each ingestion function using partial
        self.ingest_document = partial(initialize_ingest_chain_document, vector_store=vector_store)
        self.ingest_url = partial(initialize_ingest_chain_url, vector_store=vector_store)

async def initialize_ingest_chain(vector_store) -> IngestionChainWrapper:
    """
    Initializes and returns an IngestionChainWrapper with the provided vector_store.
    """
    return IngestionChainWrapper(vector_store)
